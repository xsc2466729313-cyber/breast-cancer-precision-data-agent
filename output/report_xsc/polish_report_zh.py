# -*- coding: utf-8 -*-
"""Polish submission wording and refresh Chinese architecture figures in the Word report."""
from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

from fill_original_template import FONT, set_cell, set_run_font

ROOT = Path(__file__).resolve().parent
SRC = ROOT / "徐士诚_方向1A_P5-P18_报告.docx"
ARCH = ROOT / "13_architecture.png"
LOOP = ROOT / "14_agent_loop.png"
OUT_A = Path(r"C:\Users\xsc\OneDrive\Desktop\徐士诚_方向1A_P5-P18_报告_架构图版.docx")
OUT_B = Path(r"C:\Users\xsc\OneDrive\Desktop\徐士诚_方向1A_模板填写_成员D.docx")
OUT_C = Path(r"C:\Users\xsc\OneDrive\Desktop\徐士诚_方向1A_P5-P18_报告.docx")

A_BLIP = "{http://schemas.openxmlformats.org/drawingml/2006/main}blip"
R_EMBED = "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"


def replace_paragraph_text(paragraph, new_text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        run = paragraph.add_run(new_text)
        set_run_font(run)


def fill_after_colon(paragraph, answer: str) -> None:
    text = paragraph.text.strip()
    if "：" in text:
        question = text.split("：", 1)[0] + "："
        replace_paragraph_text(paragraph, question + answer)
    else:
        replace_paragraph_text(paragraph, text + answer)


def find_para(doc, *prefixes: str):
    for prefix in prefixes:
        for p in doc.paragraphs:
            if p.text.strip().startswith(prefix):
                return p
    raise KeyError(prefixes)


def replace_image_before_caption(doc, caption_prefix: str, image_path: Path) -> bool:
    paras = list(doc.paragraphs)
    for index, para in enumerate(paras):
        if not para.text.strip().startswith(caption_prefix):
            continue
        if index == 0:
            continue
        image_para = paras[index - 1]
        blips = image_para._element.findall(f".//{A_BLIP}")
        if not blips:
            continue
        embed = blips[0].get(R_EMBED)
        if not embed:
            continue
        rel = image_para.part.rels[embed]
        rel.target_part._blob = image_path.read_bytes()
        return True
    return False


def main() -> None:
    doc = Document(str(SRC))

    replace_image_before_caption(doc, "附图（P6）：系统三层架构框图", ARCH)
    replace_image_before_caption(doc, "附图（P6）：Agent 换方法闭环", LOOP)
    replace_image_before_caption(doc, "附图（P17）：第二版换方法闭环", LOOP)
    replace_image_before_caption(doc, "附图（P17）：最终模型换方法闭环", LOOP)

    fill_after_colon(
        find_para(doc, "团队重点评价哪些方面，为什么"),
        "重点评价方向1A五项能力：找得到来源、解析得成字段、整合后可追溯、质量门能拦住错误、缺口出现后能换方法。代表科研问题只是可替换输入，不把答对某一道乳腺癌题当作评测目标。",
    )
    fill_after_colon(
        find_para(doc, "抽样数量、抽样方式和评价主体"),
        "每次真实任务对当次全量来源与全量表做规则判定；评价主体为确定性质量门和任务级诊断（来源审计、结局完整、字段完整、请求要素覆盖、科研探索可用性），与本工作台评测区同一口径，不是模型自行打分。冻结 SDTI 因 Gold Set 未加载保持未评测。",
    )
    fill_after_colon(
        find_para(doc, "正确、错误、缺失、冲突和无法判断分别如何定义"),
        "PASS：有官方来源，字段可核对，身份在同一研究内对齐，结局与问题同域。REJECT：缺 URL 或 source_id，或存在不可发布冲突。缺失：必需变量不足；建议协变量原研究未发布标为「本队列未提供」，不记系统失败。冲突：高权威来源不自动选边。REVIEW：证据不足、结局不匹配、Gold Set 未评。",
    )
    fill_after_colon(
        find_para(doc, "修正前后是否采用同一评价口径"),
        "是。质量门和任务级诊断在检索补全前后使用同一套规则。缺口补全只允许再查公开库，不允许用另一项研究的患者字段填补当前患者。",
    )

    p6_loop = find_para(doc, "本作品不是一次性检索", "本作品形成完整数据闭环")
    replace_paragraph_text(
        p6_loop,
        "本作品面向赛道二方向1A：科研问题是可替换输入，系统完成多源查找、解析、整合、溯源和修正。千问或确定性规划生成研究规格并选择工具；Adapter 调用公开库；同研究内对齐后过四层质量门。质量门失败则观察缺口并换方法，默认 8 轮、上限 12 轮。人机分工明确：研究者提供问题和 REVIEW 决策，智能体负责规划与换方法，规则层执行医学安全和发布准入。内部对照证明闭环有效：结局不同域的宽表会被拦住，半成品队列会继续换到同患者完整变量包，而不会跨库贴突变或协变量。",
    )

    fill_after_colon(
        find_para(doc, "该设计对数据结果带来的实际变化"),
        "同一输入下，仅规划得到方案；真实检索得到可审计宽表；质量门拦住结局不同域的 METABRIC；换方法后能形成治疗响应队列，并选择同患者同时含暴露与结局的完整包。任务级诊断与本工作台评测区同步，用于消融对照而不是发布 SDTI。",
    )
    fill_after_colon(
        find_para(doc, "哪些问题触发重新查找来源"),
        "必需变量覆盖不足、结局同域失败、质量门未 PASS，且仍有公开库动作可执行时触发换方法：切换独立队列、检索 GEO 目录、从文献收割研究编号，直到完整变量包或合法方法用尽。",
    )
    fill_after_colon(
        find_para(doc, "第二版相较第一版实际改善了什么"),
        "从「有表但不能分析」变成「能按问题换队列并保住来源与医学边界」。第一版 METABRIC 分析集为 0；后续补上同域治疗响应，并进一步选择同患者 PIK3CA 与响应完整包。这是查找-解析-整合-修正能力，不是只服务这一道题。",
    )
    fill_after_colon(
        find_para(doc, "哪些问题没有解决或出现了新的代价"),
        "公开队列若未发布年龄、分期、ER、PR，系统不会从其他研究贴值。代价是建议协变量可能不全；收益是结果可追溯，不会把不同研究的患者拼成假队列。宽表中的破折号保留为真实未观测。",
    )
    fill_after_colon(
        find_para(doc, "第二版数据能够支持什么后续分析"),
        "支持同患者暴露与治疗响应的描述、分组比较、来源审计和字段字典复用。不能把其他研究的年龄或受体状态当作当前患者协变量，也不能把细胞系药敏当作患者疗效。",
    )
    fill_after_colon(
        find_para(doc, "第二版是否达到团队设定的质量要求，依据是什么"),
        "已达到方向1A对可查找、可解析、可整合、可追溯、可修正的实现要求。正式 SDTI 仍未评测，依据是 Gold Set 未加载；不把外部演示页的 proxy 分数作为提交成绩。",
    )

    rows = [
        [
            "完整系统 vs 仅规划",
            "同一科研问题作为可替换输入",
            "任务级诊断：是否形成可审计宽表、来源审计能否计算",
            "真实检索形成宽表、官方来源、质量门和诊断指标",
            "仅输出研究方案与数据源规划",
            "完整系统打通问题到数据资产；规划模式可单独验收需求理解",
        ],
        [
            "完整系统 vs 去掉质量门（宽表直出）",
            "同上",
            "结局是否与问题同域、分析集是否放行、请求要素覆盖",
            "质量门拦住不同域结局，并换方法到响应队列或同患者完整包",
            "直出分子临床宽表时字段多，但不能做该问题的结局分析",
            "质量门证明查找整合包含发现错误后的修正",
        ],
        [
            "千问智能体 vs 确定性规划 / 单源停表",
            "同上",
            "问题解析、工具选择、观察后再规划、数据源多样性",
            "千问主路径完成结构化解析与函数调用；无 Key 时确定性兜底并标明 used_qwen；可换到同患者完整包",
            "规则规划稳定但复杂换题较弱；单源停在第一张表则暴露与结局无法同患者对齐",
            "对照可重复，二者共享 Adapter、质量门与本工作台评测口径",
        ],
    ]
    for i, vals in enumerate(rows):
        for j, val in enumerate(vals):
            set_cell(doc.tables[29].rows[1 + i].cells[j], val)

    fill_after_colon(
        find_para(doc, "科研数据适用性方面，本作品实际改善了什么"),
        "适用性写成四层质量门和任务级诊断，使任意科研问题都能被验收为能否支持该分析，而不是返回相关网页或一张不问结局域的宽表。",
    )
    fill_after_colon(
        find_para(doc, "数据处理方法方面，本作品实际改善了什么"),
        "形成「问题解析 → 真实检索 → 标准化与同研究对齐 → 质量门 → 换方法」完整链；原始值保留；协变量只整合同队列真实字段。",
    )
    fill_after_colon(
        find_para(doc, "结果质量和复用方面，本作品实际改善了什么"),
        "中文字段字典、官方溯源、多格式导出，以及本工作台评测区的诊断指标和消融表，便于后续科研分析、影响力评估或证据推理。",
    )
    fill_after_colon(
        find_para(doc, "没有改善的部分及增加的成本"),
        "冻结 Gold Set 未加载，正式 SDTI 未评测，不使用页面 AI proxy 冒充成绩。公开队列未发布的患者级协变量不会被补全。代价是检索轮次和人工 REVIEW 增加；收益是结果可复核、可复用。",
    )

    for p in doc.paragraphs:
        text = p.text.strip()
        if text.startswith("附图（P17）：第二版换方法闭环") or text.startswith("附图（P17）：最终模型换方法闭环"):
            replace_paragraph_text(p, "附图（P17）：智能体换方法闭环。同一输入下从不同域宽表切换到可分析队列，并保持同研究对齐。")
        elif text.startswith("附图（P17）：响应队列上同患者"):
            replace_paragraph_text(p, "附图（P17）：同研究内对齐。分子变量与治疗响应均来自可追溯公开来源，禁止跨库贴值。")
        elif text.startswith("附图（P17）：检索审计"):
            replace_paragraph_text(p, "附图（P17）：检索审计。系统按诊断持续换方法，证明发现错误后能够修正。")
        elif text.startswith("附图（P18）：模型评价中心"):
            replace_paragraph_text(p, "附图（P18）：系统评测与消融。采用任务级诊断指标，Gold Set 空则 SDTI 未评测。")
        elif text.startswith("附图（P6）：系统三层架构框图"):
            replace_paragraph_text(p, "附图（P6）：系统三层架构框图（中文）：科研问题输入、多源融合、质量闭环。")
        elif text.startswith("附图（P6）：Agent 换方法闭环") or text.startswith("附图（P6）：智能体换方法闭环"):
            replace_paragraph_text(p, "附图（P6）：智能体换方法闭环（中文）：观察、诊断、换方法、执行、判定。")

    doc.save(str(SRC))
    shutil.copy2(SRC, OUT_A)
    shutil.copy2(SRC, OUT_B)
    try:
        shutil.copy2(SRC, OUT_C)
    except PermissionError:
        print("desktop copy locked:", OUT_C)
    print("saved", SRC)
    print("copied", OUT_A)


if __name__ == "__main__":
    main()
