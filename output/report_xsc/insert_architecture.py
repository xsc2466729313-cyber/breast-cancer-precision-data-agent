# -*- coding: utf-8 -*-
"""Insert architecture diagrams into the filled member-D report and refresh P6/P17 captions."""
from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.text.paragraph import Paragraph

ROOT = Path(r"C:\Users\xsc\OneDrive\Desktop\CODEX_项目启动包_乳腺癌精准治疗科研数据智能体")
SHOTS = ROOT / "output" / "report_xsc"
DOCS_IMG = ROOT / "docs" / "images"
SRC = SHOTS / "徐士诚_方向1A_P5-P18_报告.docx"
OUT_DESKTOP_A = Path(r"C:\Users\xsc\OneDrive\Desktop\徐士诚_方向1A_P5-P18_报告.docx")
OUT_DESKTOP_B = Path(r"C:\Users\xsc\OneDrive\Desktop\徐士诚_方向1A_模板填写_成员D.docx")
FONT = "微软雅黑"
MUTED = RGBColor(0x5B, 0x64, 0x5F)


def set_run_font(run, size=10.5, italic=False, color=None):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def replace_paragraph_text(paragraph, new_text):
    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        run = paragraph.add_run(new_text)
        set_run_font(run)


def insert_after(paragraph):
    new_p = OxmlElement("w:p")
    paragraph._element.addnext(new_p)
    return Paragraph(new_p, paragraph._parent)


def add_caption_and_image(after_para, image_path: Path, caption: str, width=Inches(6.1)):
    cap = insert_after(after_para)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(caption)
    set_run_font(run, size=9, italic=True, color=MUTED)
    img = insert_after(after_para)
    img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if image_path.exists():
        img.add_run().add_picture(str(image_path), width=width)
    else:
        run = img.add_run(f"【截图缺失】{image_path.name}")
        set_run_font(run, size=9)
    return cap


def find_para(doc, prefix: str):
    for p in doc.paragraphs:
        if p.text.strip().startswith(prefix):
            return p
    raise KeyError(prefix)


def copy_architecture_pngs() -> tuple[Path, Path]:
    arch = SHOTS / "13_architecture.png"
    loop = SHOTS / "14_agent_loop.png"
    shutil.copy2(DOCS_IMG / "agent-architecture.png", arch)
    shutil.copy2(DOCS_IMG / "agent-loop.png", loop)
    return arch, loop


def already_has_architecture(doc) -> bool:
    return any("系统三层架构框图" in (p.text or "") for p in doc.paragraphs)


def fill_after_colon(paragraph, answer: str) -> None:
    text = paragraph.text.strip()
    if "：" in text:
        question = text.split("：", 1)[0] + "："
        replace_paragraph_text(paragraph, question + answer)
    else:
        replace_paragraph_text(paragraph, text + answer)


def main() -> None:
    if not SRC.exists():
        raise FileNotFoundError(SRC)
    arch, loop = copy_architecture_pngs()
    doc = Document(str(SRC))

    p6_arch = find_para(doc, "[请插入本作品实际架构图")
    if not already_has_architecture(doc):
        add_caption_and_image(p6_arch, loop, "附图（P6）：Agent 换方法闭环（观察—诊断—换方法—执行—判定）")
        add_caption_and_image(p6_arch, arch, "附图（P6）：系统三层架构框图（需求发现 → 多源融合 → 质量闭环）")

    p6_loop = None
    for prefix in (
        "[请说明本作品为什么不是一次性检索",
        "本作品不是一次性检索",
    ):
        try:
            p6_loop = find_para(doc, prefix)
            break
        except KeyError:
            continue
    if p6_loop is None:
        raise KeyError("p6_loop")
    replace_paragraph_text(
        p6_loop,
        "本作品不是一次性检索或内容摘要。质量门 REVIEW 后，Agent 观察主表、结局域与分子覆盖，诊断缺口，再检索 NCBI GEO 目录和 Europe PMC 文献，收割尚未尝试的 GSE，下载 Series Matrix 后重新对齐并过门。默认 8 轮、上限 12 轮，没有新入口则停止。代表案例：第一版 METABRIC 有 PIK3CA、无治疗响应，分析集为 0；第二版换独立队列 GSE76360，基线 50 例、48 例有治疗响应。同患者 PIK3CA 不在该 Series 中，系统不把 METABRIC 突变贴到 GEO 患者上。",
    )

    p7_change = find_para(doc, "该设计对数据结果带来的实际变化")
    fill_after_colon(
        p7_change,
        "第一版确定性规划得到 METABRIC 848×46，结局域不匹配，分析集为 0。第二版换方法解析 GSE76360，得到治疗响应队列 48 行。本机已可默认使用千问（环境变量，Key 不入库）；Function Calling 与观察后再规划已实现。不以未截取的百炼控制台冒充实测截图。",
    )

    p12_src = find_para(doc, "哪些问题触发重新查找来源")
    fill_after_colon(
        p12_src,
        "必需变量覆盖不足、结局同域失败、质量门未 PASS，且仍有公开库动作可执行（默认 8 轮、上限 12 轮未用尽）。包括换独立队列、检索 GEO 目录、从文献收割 GSE。",
    )

    p17_head = find_para(doc, "P17｜第二版输出与迭代变化")
    if not any("Agent 换方法闭环" in (p.text or "") and "P17" in (p.text or "") for p in doc.paragraphs):
        add_caption_and_image(
            p17_head,
            loop,
            "附图（P17）：第二版换方法闭环。主表从 METABRIC 空分析集切换为 GSE76360 治疗响应队列（48 行有结局）；同患者 PIK3CA 仍未补齐。",
        )

    p17_improve = find_para(doc, "第二版相较第一版实际改善了什么")
    fill_after_colon(
        p17_improve,
        "闭环在第一版结局不匹配后换方法：检索并解析 GSE76360，得到含治疗响应的独立队列（基线 50 例，48 例有结局）。分析队列从 0 变为 48 行。本轮又补上 GEO 目录检索、文献收割与千问再规划，避免只在写死清单里空转。",
    )

    doc.save(str(SRC))
    shutil.copy2(SRC, OUT_DESKTOP_A)
    shutil.copy2(SRC, OUT_DESKTOP_B)
    print(SRC)
    print(OUT_DESKTOP_A)
    print("size", SRC.stat().st_size)


if __name__ == "__main__":
    main()
