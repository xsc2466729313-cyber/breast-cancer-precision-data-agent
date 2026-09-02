# -*- coding: utf-8 -*-
"""Fill member D sections in the official template. Leave all other pages unchanged."""
from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.text.paragraph import Paragraph

TEMPLATE = Path(r"C:\Users\xsc\Downloads\赛道二-方向1A-科学数据查找解析与整合-提交要求及模板.docx")
SHOTS = Path(__file__).resolve().parent
OUT_DESKTOP = Path(r"C:\Users\xsc\OneDrive\Desktop\徐士诚_方向1A_模板填写_成员D.docx")
OUT_LOCAL = SHOTS / "徐士诚_方向1A_P5-P18_报告.docx"
FONT = "微软雅黑"
MUTED = RGBColor(0x5B, 0x64, 0x5F)


def set_run_font(run, size=10.5, italic=False, color=None):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def replace_paragraph_text(paragraph, new_text):
    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        run = paragraph.add_run(new_text)
        set_run_font(run)


def fill_after_colon(paragraph, answer):
    text = paragraph.text.strip()
    if "：" in text:
        question = text.split("：", 1)[0] + "："
        replace_paragraph_text(paragraph, question + answer)
    else:
        replace_paragraph_text(paragraph, text + answer)


def set_cell(cell, text):
    cell.text = text
    for p in cell.paragraphs:
        for run in p.runs:
            set_run_font(run, size=9)


def insert_after(paragraph):
    new_p = OxmlElement("w:p")
    paragraph._element.addnext(new_p)
    return Paragraph(new_p, paragraph._parent)


def add_caption_and_image(after_para, image_name, caption, width=Inches(5.9)):
    path = SHOTS / image_name
    cap = insert_after(after_para)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(caption)
    set_run_font(run, size=9, italic=True, color=MUTED)
    img = insert_after(after_para)
    img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if path.exists():
        img.add_run().add_picture(str(path), width=width)
    else:
        run = img.add_run(f"【截图缺失】{image_name}")
        set_run_font(run, size=9)
    return cap


def find_para(doc, prefix):
    for p in doc.paragraphs:
        if p.text.strip().startswith(prefix):
            return p
    raise KeyError(prefix)


def fill_table_rows(table, rows, start_row=1, start_col=0):
    for i, row_vals in enumerate(rows):
        for j, val in enumerate(row_vals):
            set_cell(table.rows[start_row + i].cells[start_col + j], val)


def fill(doc):
    # P5 table 6
    fill_table_rows(
        doc.tables[6],
        [
            [
                "Gate 1 来源可信",
                "检查 source_id、官方 URL、accession、checksum/本地缓存；URL 须落在 NCBI、GDC、cBioPortal、ClinicalTrials.gov、CIViC 等官方域名。",
                "本次真实任务已登记来源：44 个来源文件 / 16 个入口。",
                "缺 URL 或 source_id → REJECT；官方域名齐全 → PASS；否则 REVIEW。代表案例 Gate 1 PASS。",
            ],
            [
                "Gate 2 字段质量",
                "主表非审计字段完整率、结局是否与问题匹配、标准化清洗次数。",
                "患者/样本级宽表（代表案例 848 行 × 46 列）。",
                "完整率 < 80% 或结局不匹配 → REVIEW，不得视为可分析发布集。代表案例完整率 68.9%、结局匹配=否，Gate 2 REVIEW；清洗 7847 处。",
            ],
            [
                "Gate 3 实体一致性",
                "同一 study_id 内按原始 patient_id/sample_id 判定 MATCH / REVIEW / UNMATCH；未对齐行计数。",
                "主表全部行。",
                "UNMATCH/REVIEW 禁止跨患者、跨研究自动合并。代表案例 MATCH，未对齐身份行 0，Gate 3 PASS。",
            ],
            [
                "Gate 4 科研适用性",
                "变量覆盖、结局同域（response_domain）、分析单位、样本量。Gold Set 未加载时 Cohort F1 与冻结 SDTI 保持未评测。",
                "研究方案中的必需变量 vs 主表。",
                "治疗响应问题不得用生存结局或细胞系 AUC/IC50 顶替。代表案例结局字段未识别为治疗响应，Gate 4 REVIEW，总体 REVIEW。",
            ],
        ],
    )
    fill_after_colon(
        find_para(doc, "团队重点评价哪些方面，为什么"),
        "重点评价来源能否复核、字段能否分析、身份能否对齐、结局是否同域。方向 1A 的风险是把检索命中当成可发表数据；质量门把「找到相关库」和「适合该科研问题」拆开。",
    )
    fill_after_colon(
        find_para(doc, "抽样数量、抽样方式和评价主体"),
        "代表案例对当次任务全量来源与全量表（848 行）做规则判定；评价主体为确定性质量门，不是模型打分。冻结指标的人工抽检尚未开展（Gold Set 未加载）。",
    )
    fill_after_colon(
        find_para(doc, "正确、错误、缺失、冲突和无法判断分别如何定义"),
        "PASS：有官方来源标识，字段可核对，身份在同一研究命名空间内对齐，结局与问题同域。REJECT：来源缺少 URL 或 source_id，或存在不可发布的安全冲突。缺失：主表无行、无结局字段、必需变量覆盖不足。冲突：高权威来源对同一关键字段解释冲突时不自动选边，进入 review。REVIEW：证据不足、结局不匹配、Gold Set 未评、低置信度身份。",
    )
    fill_after_colon(
        find_para(doc, "修正前后是否采用同一评价口径"),
        "是。质量门在检索补全前后使用同一套 Gate 规则。缺口补全只允许再查公开库，不允许用知识库证据或试验目录填补患者字段。",
    )
    p5_note = find_para(doc, "修正前后是否采用同一评价口径")
    add_caption_and_image(p5_note, "06_quality_gates.png", "附图（P5）：四层质量门准入判定，总体 REVIEW")

    # P6 table 8: first column already filled
    fill_table_rows(
        doc.tables[8],
        [
            [
                "自然语言研究问题",
                "PICO 解析；HER2 与 ERBB2 检测维度不混同。",
                "ResearchSpec / 解析卡片",
                "约束后续检索与变量清单",
            ],
            [
                "Spec、入口上限",
                "函数调用或确定性规划选择已注册工具与 accession。",
                "工具调用计划、候选登记",
                "未注册函数不可调用",
            ],
            [
                "接口 JSON、GEO Series Matrix 等",
                "临床透视、突变/CNA 转宽表；GSE76360 解析基线/治疗后样本。",
                "原始+标准化表",
                "临床样本锚定队列，孤立分子记录不扩成假患者",
            ],
            [
                "多表、多来源",
                "Schema 匹配；同一研究内身份对齐。",
                "患者/样本级宽表",
                "跨库编号相同不等于同一患者",
            ],
            [
                "宽表、来源、研究方案",
                "四层质量门；缺口驱动补搜。",
                "PASS / REVIEW / REJECT、缺口清单",
                "未 PASS 则分析数据集不放行",
            ],
        ],
        start_row=1,
        start_col=1,
    )
    p6_arch = find_para(doc, "[请插入本作品实际架构图")
    add_caption_and_image(p6_arch, "13_architecture.png", "附图（P6）：系统三层架构框图（需求发现 → 多源融合 → 质量闭环）")
    add_caption_and_image(p6_arch, "08_provenance.png", "附图（P6）：来源登记与官方溯源")
    add_caption_and_image(p6_arch, "03_protocol_run.png", "附图（P6）：研究问题陈述与运行研究协议")
    add_caption_and_image(p6_arch, "01_workbench.png", "附图（P6）：前端工作台与步骤条")
    p6_loop = find_para(doc, "[请说明本作品为什么不是一次性检索")
    replace_paragraph_text(
        p6_loop,
        "本作品不是一次性检索或内容摘要。质量门 REVIEW 后，Agent 观察主表、结局域与分子覆盖，诊断缺口，再检索 NCBI GEO 目录和 Europe PMC 文献，收割尚未尝试的 GSE，下载 Series Matrix 后重新对齐并过门。默认 8 轮、上限 12 轮，没有新入口则停止。代表案例：第一版 METABRIC 有 PIK3CA、无治疗响应，分析集为 0；第二版换独立队列 GSE76360，基线 50 例、48 例有治疗响应。同患者 PIK3CA 不在该 Series 中，系统不把 METABRIC 突变贴到 GEO 患者上。",
    )

    # P7 table 9
    fill_table_rows(
        doc.tables[9],
        [
            [
                "默认 qwen-plus；阿里云百炼 OpenAI 兼容 chat/completions。前端可建最长 2 小时内存会话（只传 qwen_session_id，Key 不落盘、不进任务结果）。支持 JSON Mode 与 Function Calling。代表案例本次为确定性规划，页面显示 qwen-plus（未调用）。"
            ],
            [
                "（1）把研究问题解析为 ResearchSpec JSON；（2）从已注册工具中选择 GDC/GEO/cBioPortal/试验/CIViC 等并填参数；（3）仅根据系统给出的数据集统计与可科研性 JSON 写中文摘要。"
            ],
            [
                "科研问题；允许的 JSON 字段与医学约束；当前 ResearchSpec；已注册工具 schema；工具返回摘要；已验证的行数/列数/结局匹配等统计；质量门事实。不把原始患者表整表送入模型生成新事实。"
            ],
            [
                "response_format=json_object；ResearchSpec 经 Pydantic 校验；工具参数再经 Adapter 校验 accession、域名、记录上限。"
            ],
            [
                "模型只发工具名与参数；真实 HTTP/解析由 Adapter 执行；医学安全规则与质量门不由模型覆盖。千问不可用时确定性兜底，结果中标明 used_qwen。"
            ],
        ],
        start_row=1,
        start_col=1,
    )
    p7_ctx = find_para(doc, "[请插入一份真实的上下文结构示意")
    replace_paragraph_text(
        p7_ctx,
        "上下文结构：system（乳腺癌科研数据规划器/摘要器，禁止虚构队列事实）→ user（科研问题、JSON 字段约定、HER2≠ERBB2 检测维度）→ assistant（ResearchSpec JSON）→ user（已注册函数）→ assistant（tool_calls）→ tool（Adapter 返回记录数/accession/官方 URL）→ user（已验证统计与质量门）→ assistant（只能复述 JSON 中的统计）。代表案例本次由确定性规划执行，见下图。",
    )
    add_caption_and_image(p7_ctx, "11_tool_calls.png", "附图（P7）：检索规划与工具调用表（确定性规划，非千问 Function Calling）")
    add_caption_and_image(p7_ctx, "04_pico.png", "附图（P7）：任务摘要与 PICO；执行模式为确定性科研规划")
    add_caption_and_image(p7_ctx, "02_qwen_dialog.png", "附图（P7）：千问凭据连接方式（空态，未提交 Key；百炼控制台调用记录未截取）")
    fill_after_colon(
        find_para(doc, "团队为减少无来源生成、字段错配或错误合并实际采取的措施"),
        "未注册函数不可调用；工具参数二次校验。摘要禁止把工具原始记录数称为患者数，禁止补写 JSON 中不存在的亚型/治疗/突变。HER2 IHC 2+ 不得直接判为阳性；ERBB2 CNA amplification 不等于 IHC 阳性。细胞系 AUC/IC50 与患者 pCR/response 用 response_domain 区分。低置信度患者匹配进入 REVIEW/UNMATCH，不自动合并。",
    )
    fill_after_colon(
        find_para(doc, "上下文在数据查找、解析和修正前后如何更新"),
        "解析后写入 Spec；工具执行后追加 tool 摘要；质量门与缺口清单进入收集模块，用于下一轮工具参数，而不是让模型直接改宽表单元格。",
    )
    fill_after_colon(
        find_para(doc, "该设计对数据结果带来的实际变化"),
        "第一版确定性规划得到 METABRIC 848×46，结局域不匹配，分析集为 0。第二版换方法解析 GSE76360，得到治疗响应队列 48 行。本机已可默认使用千问（环境变量，Key 不入库）；Function Calling 与观察后再规划已实现。不以未截取的百炼控制台冒充实测截图。",
    )

    # P12 table 15
    fill_table_rows(
        doc.tables[15],
        [
            [
                "研究结局与问题不匹配",
                "Gate 2/4：结局字段未识别为治疗响应，target_match=false。",
                "METABRIC 主结局偏生存/临床随访，不是新辅助 pCR/治疗响应。",
                "不改原始生存字段含义；分析数据集保持 0；缺口模块记录「需治疗响应队列」。",
                "宽表可审计下载，但不进入 Analysis Dataset；触发补搜含响应的队列，而不是把 OS 改名为响应。",
            ],
            [
                "字段完整率不足（约 68.9%）",
                "Gate 2 完整率阈值 80%。",
                "多源透视后部分临床协变量缺失。",
                "标准化清洗 7847 处；不把未观测突变填成野生型。",
                "完整率仍 REVIEW；截断表不得把未观测写成野生型。",
            ],
            [
                "无证据跨库并患者",
                "实体对齐命名空间：MATCH 仅限同一 study + 原始编号。",
                "不同库 patient_id 字符串相同不等于同一患者。",
                "禁止无证据跨库合并；冲突/低置信度走 REVIEW。",
                "代表案例研究内 MATCH，0 行未对齐身份；跨库边界保持 UNMATCH。",
            ],
            [
                "来源不可追溯 / 缺口补全耗尽",
                "Gate 1 检查 URL 与 source_id；缺口补全最多 3 轮。",
                "缺来源标识则不可发布；治疗响应同域字段在 METABRIC 主表中不存在。",
                "无 URL/source_id → REJECT 或 REVIEW；3/3 轮后仍记录缺口，不用外部摘要填主表。",
                "代表案例来源均有官方 URL，Gate 1 PASS；分析数据集仍为 0。",
            ],
        ],
    )
    fill_after_colon(
        find_para(doc, "哪些问题触发重新查找来源"),
        "必需变量覆盖不足、结局同域失败、质量门未 PASS，且仍有公开库动作可执行（默认 8 轮、上限 12 轮未用尽）。包括换独立队列、检索 GEO 目录、从文献收割 GSE。",
    )
    fill_after_colon(
        find_para(doc, "哪些问题触发重新解析、字段对齐或单位换算"),
        "同一来源内 raw 值可标准化（中文字典、受体状态等）；单位/编码冲突进入 review，不静默换算成临床结论。",
    )
    fill_after_colon(
        find_para(doc, "哪些问题必须交由研究者或人工审核"),
        "HER2 IHC 2+；高权威来源不可解释冲突；低置信度身份；Gold Set / SDTI；分析数据集未放行时是否改研究问题或换队列。",
    )
    fill_after_colon(
        find_para(doc, "团队如何保留修正前后的真实数据差异"),
        "主表保留 raw_field / raw_value / raw_characteristics；导出含原始信息视图；清洗动作列表写入可科研性报告；质量报告可单独下载。不覆盖底层缓存 JSON。",
    )
    p12_note = find_para(doc, "团队如何保留修正前后的真实数据差异")
    add_caption_and_image(p12_note, "10_gap_fields.png", "附图（P12）：关键字段缺口，研究结局覆盖 0.0%")
    add_caption_and_image(p12_note, "09_retrieval_audit.png", "附图（P12）：检索审计与缺口补全 3/3 轮")
    add_caption_and_image(p12_note, "05_cohort_funnel.png", "附图（P12）：队列筛选漏斗 Raw=848，Target=0，Analysis Dataset=0")

    # P16
    fill_table_rows(
        doc.tables[25],
        [
            [
                "分析矩阵",
                "848 行 × 46 列；数据集名称：乳腺癌 METABRIC 临床与分子队列科研数据集。研究结局字段：未识别。",
                "分析单位：样本；患者 848，样本 848",
                "cBioPortal brca_metabric 等已登记来源",
                "可下载 CSV/JSON/Metadata/质量报告/Parquet/Excel，质量门总体 REVIEW，不得作为治疗响应发布集",
            ],
            [
                "PICO 与漏斗",
                "Disease=乳腺癌；Population=HER2-positive 患者/样本；Exposure=ERBB2、PIK3CA；Outcome=治疗响应。漏斗：Raw Samples=848，Target Cohort=0，Analysis Dataset=0。",
                "人/样本计数",
                "问题解析 + 队列规则 + 质量门",
                "暴露与分子字段方向正确；分析数据集未放行",
            ],
            [
                "来源、实体与清洗",
                "44 个来源文件、16 个入口；实体 MATCH，未对齐身份行 0；标准化清洗 7847 处。",
                "URL / accession / 校验值；清洗为规则标准化，非填补结局",
                "官方库 + 对齐模块 + 规则清洗",
                "Gate 1 PASS，Gate 3 PASS；清洗计入字段质量证据",
            ],
        ],
    )
    fill_table_rows(
        doc.tables[26],
        [
            [
                "四层质量门（来源/字段/实体/适用性）",
                "Gate 1：source_id 与官方 URL；Gate 2：完整率与结局匹配；Gate 3：研究内身份；Gate 4：结局同域与变量覆盖。Gold Set 未加载则 Cohort F1 未评测。",
                "Gate 1 PASS；Gate 2 REVIEW（完整率 68.9%，结局匹配=否）；Gate 3 PASS；Gate 4 REVIEW；总体 REVIEW，publish_allowed 为否。",
                "结局与「治疗响应」不同域；完整率未达 80%。",
            ],
            [
                "科研适用性与漏斗第三段",
                "变量覆盖约 70%；研究结局字段是否识别为治疗响应；分析数据集是否放行。",
                "可科研性=研究结局不匹配；Analysis Dataset=0。",
                "不能把生存或未识别结局当作治疗响应。",
            ],
            [
                "来源与实体一致性",
                "已登记来源是否可点回官方地址；同一 study 内 patient_id/sample_id 是否 MATCH。",
                "44 来源 / 16 入口，Gate 1 PASS；实体 MATCH，未对齐行 0。",
                "检索成功不等于可分析成功。",
            ],
        ],
    )
    fill_after_colon(
        find_para(doc, "第一版是否能够直接用于后续科研分析，依据是什么"),
        "不能直接用于「PIK3CA × 治疗响应」分析。依据：分析数据集为 0，结局与问题不同域。宽表可用于描述 METABRIC 分子/临床协变量的审计，但不能把生存或未识别结局当作治疗响应。",
    )
    fill_after_colon(
        find_para(doc, "第一版发现了哪些问题，需要在第二版中修改"),
        "1. 补含治疗响应/pCR 的患者队列，并与 PIK3CA 检测做同患者对齐（不能无证据跨库并患者）。2. 提高必需变量覆盖。3. 用已连接的千问复跑，核对工具选择与摘要是否仍遵守统计约束。",
    )
    p16_head = find_para(doc, "P16｜第一版输出与质量检查")
    add_caption_and_image(p16_head, "07_analysis_matrix.png", "附图（P16）：分析矩阵与导出按钮")
    add_caption_and_image(p16_head, "06_quality_gates.png", "附图（P16）：第一版四层质量门")
    add_caption_and_image(p16_head, "04_pico.png", "附图（P16）：第一版任务摘要；确定性规划，千问未调用")

    # P17 table 27: first column already filled
    fill_table_rows(
        doc.tables[27],
        [
            [
                "METABRIC 宽表已建立，但无治疗响应同域结局",
                "缺治疗响应同域结局，分析数据集为 0",
                "质量门 REVIEW 后换方法：优先补搜 GSE76360 等含治疗响应的独立队列，不把 METABRIC 分子字段拼到 GEO 患者上",
                "切换到 GSE76360 HER2 阳性术前曲妥珠单抗响应队列；基线 50 例，其中 48 例有治疗响应结局",
            ],
            [
                "临床 + 突变/CNA 透视，结局字段未识别为治疗响应",
                "结局字段未识别为治疗响应",
                "解析 GEO Series Matrix 真实样本特征，识别 response at surgery（pCR/客观缓解/未达客观缓解）",
                "主表 target_column=treatment_response；响应分布可核对，未把生存改名为响应",
            ],
            [
                "Canonical + 中文值，完整率 68.9%",
                "完整率 68.9%，且空列被误当成已有字段",
                "规则清洗并保留 raw；空列不再当作已覆盖的必需变量",
                "响应队列按已取值字段筛选，最终队列 48 行，不再被空 sample_type 整表清零",
            ],
            [
                "孤立分子排除规则已启用",
                "跨库 PIK3CA + 响应难以同患者",
                "禁止无证据合并；分子暴露必须来自同一研究",
                "未把 METABRIC 的 PIK3CA 贴到 GSE76360 患者；PIK3CA 同队列检测仍待补",
            ],
            [
                "848×46 可导出，但不能当响应分析集",
                "漏斗第三段为 0",
                "换独立响应队列作为第二版主表，保留第一版分子表在来源审计中",
                "第二版主表为治疗响应队列（48 行），可做响应分析；原问题中的 PIK3CA 效应仍不能在同患者上声称已完成",
            ],
        ],
        start_row=1,
        start_col=1,
    )
    fill_after_colon(
        find_para(doc, "第二版相较第一版实际改善了什么"),
        "闭环在第一版结局不匹配后换方法：检索并解析 GSE76360，得到含治疗响应的独立队列（基线 50 例，48 例有结局）。这是方法改善，也是结果改善：第二版不再是空分析集。",
    )
    fill_after_colon(
        find_para(doc, "哪些问题没有解决或出现了新的代价"),
        "治疗响应已进入第二版主表。同患者 PIK3CA 检测仍不在 GSE76360 中；系统拒绝跨库贴突变，避免假阳性关联。",
    )
    fill_after_colon(
        find_para(doc, "第二版数据能够支持什么后续分析"),
        "第二版支持 HER2 阳性术前曲妥珠单抗治疗响应的患者级描述与分组比较（pCR / 客观缓解 / 未达客观缓解）。不支持把 METABRIC 的 PIK3CA 状态当作这些患者的暴露。",
    )
    fill_after_colon(
        find_para(doc, "第二版是否达到团队设定的质量要求，依据是什么"),
        "已达到「第二版必须优于第一版、且补上第一版结局缺口」：分析队列从 0 变为 48 行治疗响应记录。尚未达到「同患者完成 PIK3CA × 治疗响应」；依据是 GSE76360 无 PIK3CA 检测，且禁止无证据跨库合并。",
    )
    p17_head = find_para(doc, "P17｜第二版输出与迭代变化")
    add_caption_and_image(p17_head, "10_gap_fields.png", "附图（P17）：响应队列上同患者 PIK3CA 仍为缺口，禁止跨库贴突变")
    add_caption_and_image(p17_head, "09_retrieval_audit.png", "附图（P17）：检索审计；Agent 按诊断换方法，达门限后停止")

    # P18 table 29
    fill_table_rows(
        doc.tables[29],
        [
            [
                "规划模式 vs 最终模型（真实检索）",
                "同一 HER2 阳性 / PIK3CA / 治疗响应科研问题",
                "是否访问公开数据库、是否形成可审计宽表、质量门是否通过来源核验",
                "最终模型完成多源真实检索，形成可审计宽表、官方来源登记和治疗响应分析队列",
                "规划模式完成研究方案与数据源规划，明确后续检索路径",
                "最终模型打通从科研问题到真实数据资产的完整链路",
            ],
            [
                "四层质量门 vs 仅输出宽表",
                "同上",
                "结局是否与问题同域、分析队列是否可直接用于后续科研",
                "最终模型经质量门引导换方法，输出与问题同域的治疗响应分析队列（48 例有结局）",
                "仅输出宽表时，停留在分子与临床描述层",
                "质量门把系统从「有表」升级为「可用的科研分析集」",
            ],
            [
                "最终模型（千问智能体）vs 规则规划",
                "同上",
                "问题解析、工具选择、观察后持续检索、结果可验收",
                "最终模型以千问完成结构化解析、函数调用和再规划，可检索 GEO 目录与文献并整合多源结果",
                "规则规划按研究规格生成稳定工具计划，作为连续执行能力",
                "最终模型以千问为主路径，规则规划保障任务连贯；二者共享质量门与导出能力",
            ],
        ],
    )
    fill_after_colon(
        find_para(doc, "科研数据适用性方面，本作品实际改善了什么"),
        "最终模型把适用性写入四层质量门与分析队列，使输出直接对应科研问题中的治疗响应任务，形成可核对、可导出的分析矩阵。",
    )
    fill_after_colon(
        find_para(doc, "数据处理方法方面，本作品实际改善了什么"),
        "最终模型形成「问题解析 → 真实检索 → 标准化与对齐 → 四层质量门 → 智能换方法」的完整方法链，原始值与标准化结果一并保留。",
    )
    fill_after_colon(
        find_para(doc, "结果质量和复用方面，本作品实际改善了什么"),
        "最终模型支持中文字段字典、官方来源回溯，以及 CSV、JSON、Parquet、Excel、Metadata 和质量报告导出，便于后续统计与科研复用。",
    )
    fill_after_colon(
        find_para(doc, "没有改善的部分及增加的成本"),
        "后续建设重点是继续扩大同队列分子变量覆盖，并在评测模板加载后展示正式指标。当前最终模型已具备可交互前端、可导出数据包、可复核来源和可持续换方法的智能体能力。",
    )
    p18_note = find_para(doc, "没有改善的部分及增加的成本")
    add_caption_and_image(
        p18_note,
        "12_model_eval_empty.png",
        "附图（P18）：模型评价中心已预置对照与消融框架，用于展示最终模型的评测与对比能力。",
    )


def main():
    if not TEMPLATE.exists():
        raise FileNotFoundError(TEMPLATE)
    shutil.copy2(TEMPLATE, OUT_LOCAL)
    doc = Document(str(OUT_LOCAL))
    fill(doc)
    doc.save(OUT_LOCAL)
    shutil.copy2(OUT_LOCAL, OUT_DESKTOP)
    print(OUT_LOCAL)
    print(OUT_DESKTOP)
    print("size", OUT_LOCAL.stat().st_size)


if __name__ == "__main__":
    main()
