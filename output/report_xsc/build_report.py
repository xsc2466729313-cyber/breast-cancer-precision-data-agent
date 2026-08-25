# -*- coding: utf-8 -*-
"""Build Xu Shicheng direction-1A Word report with embedded screenshots."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent
DESKTOP = Path(r"C:\Users\xsc\OneDrive\Desktop")
FONT = "微软雅黑"
ACCENT = RGBColor(0x1F, 0x4E, 0x46)
MUTED = RGBColor(0x5B, 0x64, 0x5F)


def set_run_font(run, size=11, bold=False, color=None, italic=False):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def add_para(doc, text, size=11, bold=False, space_after=8, color=None, align=None, first_line=True):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.28
    if first_line and align is None:
        pf.first_line_indent = Cm(0.74)
    if align is not None:
        p.alignment = align
        pf.first_line_indent = Cm(0)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    return p


def add_heading_cn(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(16 if level == 1 else 12)
    p.paragraph_format.space_after = Pt(8)
    for run in p.runs:
        set_run_font(run, size=16 if level == 1 else 13 if level == 2 else 12, bold=True, color=ACCENT)
    return p


def shade_header(cell, fill="1F4E46"):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_run_font(run, size=9, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        shade_header(cell)
    for r_i, row in enumerate(rows, start=1):
        for c_i, val in enumerate(row):
            cell = table.rows[r_i].cells[c_i]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            set_run_font(run, size=9)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table


def add_shot(doc, filename, caption):
    path = ROOT / filename
    if not path.exists():
        add_para(doc, f"【截图缺失】{filename}", first_line=False, color=RGBColor(0xA3, 0x3B, 0x24))
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run()
    run.add_picture(str(path), width=Inches(6.15))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(12)
    cap.paragraph_format.first_line_indent = Cm(0)
    r = cap.add_run(caption)
    set_run_font(r, size=9, italic=True, color=MUTED)


def configure_styles(doc):
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    for i in range(1, 4):
        style = doc.styles[f"Heading {i}"]
        style.font.name = FONT
        style.font.color.rgb = ACCENT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)


def build():
    doc = Document()
    configure_styles(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(36)
    r = title.add_run("赛道二 · 方向 1A")
    set_run_font(r, size=14, color=ACCENT)

    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = h.add_run("科学数据查找、解析与整合")
    set_run_font(r, size=22, bold=True, color=ACCENT)

    h2 = doc.add_paragraph()
    h2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = h2.add_run("成员 D 报告稿（P5 / P6 / P7 / P12 / P16 / P17 / P18）")
    set_run_font(r, size=14, bold=True)

    meta_lines = [
        "报告人：徐士诚（成员 D）",
        "作品：乳腺癌精准治疗科研数据智能体  v2.0.0-qwen-agent",
        "代表案例：研究 HER2 阳性乳腺癌中 PIK3CA 突变是否影响治疗响应，并整理患者级科研数据集",
        "实测日期：2026-08-25",
        "复现入口：http://127.0.0.1:8002/",
    ]
    for line in meta_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(line)
        set_run_font(r, size=11)

    add_para(
        doc,
        "说明：本稿只写本仓库已实现、并在本机真实数据模式下跑通的内容。Gold Set 未加载，因此 Retrieval Precision/Recall、Faithfulness、SDTI 等正式成绩一律写「未评测」，不编造分数。截图由本机浏览器对本地服务直接截取，文档中不含 API Key。",
        first_line=True,
        space_after=14,
    )

    add_heading_cn(doc, "0. 实测结论与材料边界", 1)
    add_heading_cn(doc, "0.1 两次已跑通任务", 2)
    add_table(
        doc,
        ["执行模式", "分析矩阵", "质量门总体", "实体对齐", "说明"],
        [
            ["仅生成检索计划", "0 行 × 0 列", "REVIEW", "UNMATCH", "生成研究方案与数据源规划，不访问外部库"],
            [
                "真实数据模式（本稿主证据）",
                "848 行 × 46 列",
                "REVIEW",
                "MATCH（同一研究内可对齐）",
                "cBioPortal METABRIC 临床与分子队列；结局按「治疗响应」核对后不匹配，分析数据集为 0",
            ],
        ],
    )
    add_para(
        doc,
        "本次为生成本稿再次运行真实数据模式。页面显示：任务完成；执行模式为「确定性科研规划」；推理模型「qwen-plus（未调用）」；分析矩阵 848 行 × 46 列。来源 Gate 1 PASS，实体 Gate 3 PASS，字段质量与科研适用性 REVIEW，漏斗 Raw Samples=848、Target Cohort=0、Analysis Dataset=0。",
    )

    add_heading_cn(doc, "0.2 本稿能证明什么、不能证明什么", 2)
    add_table(
        doc,
        ["材料", "状态", "用途"],
        [
            ["前端工作台、研究问题陈述、运行控件", "已截图", "P6 架构"],
            ["千问凭据对话框（空 Key，仅接口地址）", "已截图", "P7 调用方式，不是已启用会话"],
            ["PICO、质量门四卡、漏斗、分析矩阵、溯源、检索审计、工具表", "已截图", "P5/P12/P16/P17"],
            ["模型评价中心空态", "已截图", "P18：标明未跑，不当成绩"],
            ["千问 Function Calling 真实调用（used_qwen=true）", "未完成", "本机未配置百炼 Key，P7 不得写「本次由千问完成工具选择」"],
            ["阿里云百炼控制台调用记录", "未完成", "无账号登录权限，不能代截"],
            ["外部 GitHub 同口径数字对照（BEIR/Valentine 等）", "未实测", "P18 只写内部变体，外部不填分数"],
        ],
    )

    add_heading_cn(doc, "1. 截图证据", 1)
    add_para(doc, "以下截图均取自本机 http://127.0.0.1:8002/ ，截图时未填写、未展示 API Key。", first_line=True)

    add_heading_cn(doc, "图 1  工作台全页", 2)
    add_para(doc, "对应 P6。可见步骤条、研究问题陈述入口与系统在线状态。配置条标明「千问未配置 / 当前使用确定性规划」，不得把该页解读为已启用千问 Agent。")
    add_shot(doc, "01_workbench.png", "图1  工作台：科研问题入口与确定性规划提示")

    add_heading_cn(doc, "图 2  研究问题陈述与运行参数", 2)
    add_para(doc, "代表案例原文、执行模式=真实数据模式、检索入口上限=8、单表记录上限=10,000、缺口补全=3 轮；启用模型规划与确定性兜底。")
    add_shot(doc, "03_protocol_run.png", "图2  研究问题陈述与真实数据模式参数")

    add_heading_cn(doc, "图 3  千问凭据连接方式（空态）", 2)
    add_para(
        doc,
        "对应 P7。系统提供手工录入或本机百炼 CSV 导入；默认兼容地址为 https://dashscope.aliyuncs.com/compatible-mode/v1，模型 qwen-plus。安全说明写明凭据仅进程内存保存、最长 2 小时、不写入数据库/日志/任务结果。本稿截取时状态为「尚未提交凭据」，因此只能证明调用链路已实现，不能证明本次任务调用了千问。",
    )
    add_shot(doc, "02_qwen_dialog.png", "图3  千问 API 配置对话框（未提交 Key）")

    add_heading_cn(doc, "图 4  任务摘要与 PICO", 2)
    add_para(doc, "对应 P7、P16。Disease=乳腺癌；Population=HER2-positive 患者/样本；Exposure=ERBB2、PIK3CA；Outcome=治疗响应。页面同时写明 848 行宽表已生成，但队列缺少匹配的治疗响应结局，未用生存数据顶替。")
    add_shot(doc, "04_pico.png", "图4  任务完成态：确定性规划 + PICO 五格")

    add_heading_cn(doc, "图 5  队列筛选漏斗", 2)
    add_para(doc, "对应 P12、P16。来源行数 848；最终队列 0；患者数/样本数在最终队列口径下为 0；变量覆盖 70.0%；患者 Linkage F1 未评测。漏斗三段为 Raw Samples=848、Target Cohort=0、Analysis Dataset=0。")
    add_shot(doc, "05_cohort_funnel.png", "图5  队列筛选路径：分析数据集未放行")

    add_heading_cn(doc, "图 6  四层质量门", 2)
    add_para(doc, "对应 P5、P12、P16。Gate 1 来源可信 PASS（44 来源，官方 URL）；Gate 2 字段质量 REVIEW（完整率 68.9%，清洗 7847 处，结局匹配=否）；Gate 3 实体一致性 PASS；Gate 4 科研适用性 REVIEW（治疗响应结局未识别）。总体 REVIEW。")
    add_shot(doc, "06_quality_gates.png", "图6  四层质量门准入判定")

    add_heading_cn(doc, "图 7  分析矩阵与导出", 2)
    add_para(doc, "对应 P16。数据集名称：乳腺癌 METABRIC 临床与分子队列科研数据集。848 患者 / 848 样本；研究结局字段：未识别。可下载 CSV / JSON / Metadata / 质量报告 / Parquet / Excel。该宽表可审计，不能当作治疗响应分析发布集。")
    add_shot(doc, "07_analysis_matrix.png", "图7  分析矩阵表头与导出按钮")

    add_heading_cn(doc, "图 8  来源登记与官方溯源", 2)
    add_para(doc, "对应 P6、P12。44 个来源文件、16 个入口；主路径为 cBioPortal（24 项），并登记 NCBI GEO、GDC、CIViC、ClinicalTrials.gov。每条来源可回到官方地址。")
    add_shot(doc, "08_provenance.png", "图8  多源检索到主科研数据集的溯源图")

    add_heading_cn(doc, "图 9  检索审计与缺口补全", 2)
    add_para(doc, "对应 P12、P17。缺口补全 3/3 轮已用尽，状态仍为需要继续补全；关键缺口包括疾病字段覆盖与治疗响应同域结局。系统记录补救方向，不用外部摘要或其他患者填补主表。")
    add_shot(doc, "09_retrieval_audit.png", "图9  检索审计：缺口补全轮次与来源表")
    add_shot(doc, "10_gap_fields.png", "图9续  关键字段缺口：研究结局覆盖 0.0%")

    add_heading_cn(doc, "图 10  工具调用表", 2)
    add_para(doc, "对应 P7、P16。8/8 工具成功；登记 44 个真实来源。可见 cBioPortal（METABRIC / TCGA）、GEO GSE76360、GDC、CIViC、ClinicalTrials.gov 的真实接口调用记录。本次执行模式仍为确定性规划，不是千问 Function Calling。")
    add_shot(doc, "11_tool_calls.png", "图10  检索规划与工具调用完成表")

    add_heading_cn(doc, "图 11  模型评价中心（空态，非成绩）", 2)
    add_para(doc, "对应 P18。页面写明独立于科研数据搜寻闭环；测试报告未生成；真实观测行 0；已连接模型 0；Gold Set 正式成绩未评测。此图只证明评价页存在且拒绝空成绩，不能当作模型对比结果。")
    add_shot(doc, "12_model_eval_empty.png", "图11  模型评价中心空态，不当作正式成绩")

    add_heading_cn(doc, "P5  本作品采用的数据质量评价方法", 1)
    add_heading_cn(doc, "实际评价方案", 2)
    add_para(
        doc,
        "本作品把质量评价做成发布准入，而不是演示用综合分。评价分四层，全部基于本次任务观察到的来源、宽表、身份对齐和科研问题匹配。没有 Gold Set 时，Cohort F1 与冻结 SDTI 保持未评测。",
    )
    add_heading_cn(doc, "实际采用的评价内容", 2)
    add_table(
        doc,
        ["实际采用的评价内容", "具体评价方法", "评价对象与样本", "结果如何影响后续处理"],
        [
            [
                "Gate 1 来源可信",
                "检查 source_id、官方 URL、accession、checksum/本地缓存；URL 须落在 NCBI、GDC、cBioPortal、ClinicalTrials.gov、CIViC、Europe PMC 等官方域名",
                "本次任务已登记来源（44 个来源文件 / 16 个入口）",
                "缺 URL 或 source_id → REJECT；官方域名齐全 → PASS；否则 REVIEW",
            ],
            [
                "Gate 2 字段质量",
                "主表非审计字段完整率、结局是否与问题匹配、标准化清洗次数",
                "患者/样本级宽表（848×46）",
                "完整率 < 80% 或结局不匹配 → REVIEW，不得视为可分析发布集",
            ],
            [
                "Gate 3 实体一致性",
                "同一 study_id 内按原始 patient_id/sample_id 判定 MATCH / REVIEW / UNMATCH；未对齐行计数",
                "主表全部行",
                "UNMATCH/REVIEW 禁止跨患者、跨研究自动合并",
            ],
            [
                "Gate 4 科研适用性",
                "变量覆盖、结局同域（response_domain）、分析单位、样本量",
                "研究方案中的必需变量 vs 主表",
                "治疗响应问题不得用生存结局或细胞系 AUC/IC50 顶替",
            ],
            [
                "任务级诊断（非官方成绩）",
                "来源审计完整度、结局完整率、字段完整率、请求要素覆盖、科研探索可用性",
                "当前任务结果",
                "只用于说明适用性缺口，不写成 SDTI",
            ],
            [
                "冻结 SDTI（公式已实现、成绩未评）",
                "docs/06 评测指标与 SDTI：Retrieval F1、Faithfulness、Traceability、Error F1、Repair Accuracy 的几何平均",
                "需冻结 Gold Set",
                "当前模板为空，报告中一律写未评测",
            ],
        ],
    )
    add_heading_cn(doc, "评价口径说明", 2)
    add_para(
        doc,
        "团队重点评价「来源能否复核、字段能否分析、身份能否对齐、结局是否同域」。方向 1A 的风险是把检索命中当成可发表数据；质量门把「找到相关库」和「适合该科研问题」拆开。",
    )
    add_para(
        doc,
        "抽样数量、抽样方式和评价主体：代表案例对当次任务全量来源与全量表（848 行）做规则判定，评价主体为确定性质量门，不是模型打分。冻结指标的人工抽检尚未开展（Gold Set 未加载）。",
    )
    add_table(
        doc,
        ["状态", "定义"],
        [
            ["正确 / PASS", "有官方来源标识；字段与类型可核对；身份在同一研究命名空间内对齐；结局字段与问题同域"],
            ["错误 / REJECT", "来源缺少 URL 或 source_id；或存在不可发布的安全冲突"],
            ["缺失", "主表无行、无结局字段、必需变量覆盖不足"],
            ["冲突", "高权威来源对同一关键字段解释冲突时不自动选边，进入 review"],
            ["无法判断 / REVIEW", "证据不足、结局不匹配、Gold Set 未评、低置信度身份"],
        ],
    )
    add_para(
        doc,
        "修正前后采用同一评价口径。质量门在检索补全前后使用同一套 Gate 规则。缺口补全只允许再查公开库，不允许用知识库证据或试验目录填补患者字段。",
    )

    add_heading_cn(doc, "P6  系统总体架构与数据闭环", 1)
    add_heading_cn(doc, "本作品真实架构", 2)
    for line in [
        "研究问题陈述",
        "→ Qwen JSON Mode 或确定性解析 → PICO / ResearchSpec",
        "→ 研究方案与数据源规划",
        "→ Qwen Function Calling 或确定性工具规划",
        "→ 受控工具：GDC / GEO / cBioPortal / ClinicalTrials.gov / CIViC / BioSample / Europe PMC",
        "→ 原始表解析 → Canonical 字段 + raw_field/raw_value + source_id",
        "→ 实体对齐（study 内 MATCH/REVIEW/UNMATCH，禁止无证据跨库并患者）",
        "→ 四层质量门 PASS / REVIEW / REJECT",
        "→ REVIEW 且开启缺口补全：迭代检索（最多 3 轮）→ 再进入质量门",
        "→ 分析矩阵 + 中文字段字典 + 质量报告 + 来源溯源",
        "→ 导出 CSV / JSON / Parquet / Excel / Metadata / 质量报告",
    ]:
        add_para(doc, line, first_line=False, space_after=2)
    add_para(
        doc,
        "Qwen 位置：只在问题解析、工具选择、基于已验证统计的中文摘要。外部工具位置：公开数据库 Adapter。人工位置：REVIEW/REJECT、低置信度身份、Gold Set 与正式 SDTI。",
    )
    add_heading_cn(doc, "实际模块", 2)
    add_table(
        doc,
        ["实际模块", "输入", "本作品的核心处理", "输出", "与下一模块的关系"],
        [
            ["科研数据需求理解", "自然语言研究问题", "PICO 解析；HER2 与 ERBB2 检测维度不混同", "ResearchSpec / 解析卡片", "约束后续检索与变量清单"],
            ["数据源发现与筛选", "Spec、入口上限", "函数调用或确定性规划选择已注册工具与 accession", "工具调用计划、候选登记", "未注册函数不可调用"],
            [
                "多类型内容解析",
                "接口 JSON、GEO Series Matrix 等",
                "临床透视、突变/CNA 转宽表；GSE76360 解析基线/治疗后样本",
                "原始+标准化表",
                "临床样本锚定队列，孤立分子记录不扩成假患者",
            ],
            ["字段对齐与多源整合", "多表、多来源", "Schema 匹配；同一研究内身份对齐", "患者/样本级宽表", "跨库编号相同不等于同一患者"],
            ["质量检查与反馈修正", "宽表、来源、研究方案", "四层质量门；缺口驱动补搜", "PASS/REVIEW/REJECT、缺口清单", "未 PASS 则分析数据集不放行"],
            ["结构化输出", "通过或待审查的任务结果", "字典、溯源、导出", "可下载科研数据资产", "供统计/机器学习后续使用，不作诊疗"],
        ],
    )
    add_heading_cn(doc, "数据闭环的关键设计", 2)
    add_para(
        doc,
        "本作品不是一次性检索摘要。质量门 REVIEW 后，收集模块根据缺失的必需变量提出下一轮公开库动作，再解析、再对齐、再过门。代表案例说明闭环有效性：METABRIC 可提供 PIK3CA/ERBB2 分子变量，但不含与问题匹配的治疗响应字段。系统将分析数据集保持为 0，而不是用总生存或知识证据顶替结局。这是闭环拦截，不是检索失败后的文案掩饰。",
    )
    add_para(
        doc,
        "最终数据成果：任务级患者/样本宽表、中文字段字典、来源清单（URL/accession/checksum）、四层质量门报告、可导出 Metadata 与质量报告。真实数据模式已实际请求 cBioPortal（含 METABRIC）、NCBI GEO、GDC、ClinicalTrials.gov/AACT、CIViC。GTEx、OncoKB 等仅出现在数据源规划中，不生成虚假记录。",
    )

    add_heading_cn(doc, "P7  Qwen 使用方式与上下文工程", 1)
    add_para(
        doc,
        "诚实口径：调用链路已实现（见图3），代表案例本次由确定性规划执行（见图4「qwen-plus（未调用）」）。在连接百炼并出现 used_qwen=true 之前，不把工具选择归功于千问。",
        bold=False,
    )
    add_table(
        doc,
        ["内容", "本作品实际做法"],
        [
            [
                "使用的 Qwen 模型与调用方式",
                "默认 qwen-plus；阿里云百炼 OpenAI 兼容 chat/completions。前端可建最长 2 小时内存会话（只传 qwen_session_id，Key 不落盘、不进任务结果）。支持 JSON Mode 与 Function Calling",
            ],
            [
                "Qwen 承担的具体任务",
                "（1）把研究问题解析为 ResearchSpec JSON；（2）从已注册工具中选择 GDC/GEO/cBioPortal/试验/CIViC 等并填参数；（3）仅根据系统给出的数据集统计与可科研性 JSON 写中文摘要",
            ],
            [
                "一次查找时提供给模型的上下文",
                "科研问题；允许的 JSON 字段与医学约束；当前 ResearchSpec；已注册工具 schema；工具返回摘要；已验证的行数/列数/结局匹配等统计；质量门事实。不把原始患者表整表送入模型生成新事实",
            ],
            [
                "结构化输出或字段约束",
                "response_format=json_object；ResearchSpec 经 Pydantic 校验；工具参数再经 Adapter 校验 accession、域名、记录上限",
            ],
            [
                "与检索、解析、数据库工具的协作",
                "模型只发工具名与参数；真实 HTTP/解析由 Adapter 执行；医学安全规则与质量门不由模型覆盖",
            ],
        ],
    )
    add_para(doc, "上下文组织（结构示意）：system 规划器/摘要器禁止虚构队列事实 → user 科研问题与 JSON 字段约定 → assistant 输出 ResearchSpec JSON → user 给出已注册函数 → assistant 发出 tool_calls → tool 返回由 Adapter 产生的记录数/accession/官方 URL → user 给出已验证统计 → assistant 只能复述 JSON 中的统计。")
    add_para(
        doc,
        "减少无来源生成的措施：未注册函数不可调用；摘要禁止把工具原始记录数称为患者数；HER2 IHC 2+ 不得直接判阳；ERBB2 CNA amplification 不等于 IHC 阳性；细胞系 AUC/IC50 与患者 pCR/response 用 response_domain 区分；低置信度匹配进入 REVIEW/UNMATCH；千问不可用时确定性兜底并标明 used_qwen。",
    )

    add_heading_cn(doc, "P12  质量检查与实际反馈修正机制", 1)
    add_table(
        doc,
        ["实际质量问题", "系统如何发现", "原因分析", "实际修改的环节", "修正后的结果", "反馈如何改变数据流程"],
        [
            [
                "研究结局与问题不匹配",
                "Gate 2/4：target_match=false，结局字段未识别为治疗响应",
                "METABRIC 主结局偏生存/临床随访，不是新辅助 pCR/治疗响应",
                "不改原始生存字段含义；分析数据集保持 0；缺口模块记录需治疗响应队列",
                "宽表可审计下载，但不进入 Analysis Dataset",
                "触发补搜 GEO/含响应的队列，而不是把 OS 改名为响应",
            ],
            [
                "字段完整率不足（约 68.9%）",
                "Gate 2 完整率阈值 80%",
                "多源透视后部分临床协变量缺失",
                "标准化清洗 7847 处；不填造缺失突变为野生型",
                "完整率仍 REVIEW",
                "截断表不得把未观测写成野生型",
            ],
            [
                "治疗后配对样本泄漏风险",
                "GSE76360 解析规则",
                "同一患者 baseline/post 若按行随机切分会泄漏",
                "主表保留基线，治疗后样本分离",
                "避免错误分析单位",
                "属解析规则，不经模型改写",
            ],
            [
                "孤立分子记录",
                "临床样本锚定",
                "mutation/CNA 无法连到临床样本",
                "排除计数，不扩成新患者",
                "降低假高缺失率",
                "整合阶段硬约束",
            ],
            [
                "无证据跨库并患者",
                "实体对齐命名空间",
                "不同库 patient_id 字符串相同",
                "MATCH 仅限同一 study + 原始编号",
                "代表案例 MATCH，0 行未对齐身份",
                "冲突/低置信度走 REVIEW",
            ],
            [
                "来源不可追溯",
                "Gate 1",
                "缺 URL 或 source_id",
                "REJECT 或 REVIEW",
                "代表案例来源均有官方 URL，Gate 1 PASS",
                "无来源记录不得发布",
            ],
        ],
    )
    add_para(doc, "触发重新查找来源的问题：必需变量覆盖不足、结局同域失败、质量门未 PASS，且仍有公开库动作可执行。触发重新解析的问题：同一来源内 raw 值可标准化；单位/编码冲突进入 review，不静默换算成临床结论。必须交由研究者审核：HER2 IHC 2+、高权威来源不可解释冲突、低置信度身份、Gold Set/SDTI、分析数据集未放行时是否改问题或换队列。")
    add_para(doc, "修正前后差异的保留方式：主表保留 raw_field / raw_value / raw_characteristics；导出含原始信息视图；清洗动作列表写入可科研性报告；质量报告可单独下载。不覆盖底层缓存 JSON。")

    add_heading_cn(doc, "P16  第一版输出与质量检查", 1)
    add_para(doc, "第一版指真实数据模式首次完整输出（确定性规划，千问未调用）。若统稿需要更早版本，可将「仅规划、0 行」作为第 0 版，不与第一版数字混用。")
    add_table(
        doc,
        ["对象、记录或条目", "主要内容或字段", "单位/说明", "来源", "核对状态或问题"],
        [
            ["分析矩阵", "848 行 × 46 列；METABRIC 临床与分子队列", "分析单位：样本", "cBioPortal brca_metabric 等已登记来源", "可下载，但质量门 REVIEW"],
            ["PICO", "Disease=乳腺癌；Population=HER2-positive；Exposure=ERBB2、PIK3CA；Outcome=治疗响应", "—", "问题解析", "暴露与分子字段方向正确"],
            ["漏斗", "Raw Samples=848；Target Cohort=0；Analysis Dataset=0", "人/样本计数", "队列规则 + 质量门", "未放行分析集"],
            ["来源", "44 个来源文件，16 个入口", "URL/accession/校验值", "官方库", "Gate 1 PASS"],
            ["实体", "MATCH，未对齐身份行 0", "研究内编号", "对齐模块", "Gate 3 PASS"],
            ["清洗", "7847 处标准化", "非填补结局", "规则清洗", "计入字段质量证据"],
        ],
    )
    add_para(doc, "第一版质量检查：Gate 1 PASS；Gate 2 REVIEW（完整率 68.9%，结局匹配=否）；Gate 3 PASS；Gate 4 REVIEW（结局不匹配，变量覆盖约 70%）；总体 REVIEW，publish_allowed 为否。")
    add_para(
        doc,
        "第一版不能直接用于「PIK3CA × 治疗响应」分析。依据：分析数据集为 0，结局与问题不同域。宽表可用于描述 METABRIC 分子/临床协变量的审计，但不能把生存或未识别结局当作治疗响应。后续修改方向：补含治疗响应/pCR 的患者队列并与 PIK3CA 检测做同患者对齐（不能无证据跨库并患者）；提高必需变量覆盖；用已连接的千问复跑以核验工具选择。",
    )

    add_heading_cn(doc, "P17  第二版输出与迭代变化", 1)
    add_para(
        doc,
        "模板允许：若没有形成独立第二版发布集，如实写修改记录，不要虚构两版都变好的数字。本作品已实现闭环代码（缺口补全最多 3 轮），代表案例在第一版之后未能得到可放行的第二版分析数据集。原因不是程序中断，而是科学约束：继续检索不能把 METABRIC 生存结局改写成治疗响应，也不能把 CIViC 证据行写成患者疗效。",
    )
    add_table(
        doc,
        ["变化内容", "第一版结果", "发现的问题", "实际修改的方法", "第二版结果"],
        [
            ["数据源或查找范围", "METABRIC 宽表已建立", "缺治疗响应同域结局", "质量门 REVIEW → 缺口补全提出再查响应队列/GEO", "分析数据集仍为 0（未放行）"],
            ["内容解析", "临床+突变/CNA 透视", "结局字段未识别", "不改原始字段语义", "保持原始生存/临床字段，不改名冒充响应"],
            ["字段统一", "Canonical + 中文值", "完整率 68.9%", "规则清洗，保留 raw", "清洗已执行，完整率未达 PASS"],
            ["重复、缺失、冲突", "孤立分子排除规则已启用", "跨库 PIK3CA+响应难以同患者", "禁止无证据合并", "保持 UNMATCH 跨库边界"],
            ["结构化输出", "848×46 可导出", "不能当响应分析集", "漏斗第三段为 0", "仍为 REVIEW 资产，非发布集"],
        ],
    )
    add_para(doc, "第二版相较第一版实际改善的是方法：把「已有宽表」和「可分析」分开，缺口与工具日志可审计。这不是结局覆盖率的数值改善。未解决的问题：治疗响应仍未进入主表。补搜可能增加试验目录/证据条目，当前规则禁止用其顶替患者主表，代价是分析数据集为空。")
    add_para(doc, "第二版在未换队列前，仅支持 METABRIC 上与已有字段一致的描述性分子-临床审计，不支持原问题中的治疗响应关联分析。未达到「质量门 PASS 且分析数据集非空」。若答辩需要可见的第二版成功队列，应另跑一条结局本身就是治疗响应的任务（例如 GSE76360 HER2 基线响应队列）作为案例 B，不要把案例 A 的空分析集改写成成功。")

    add_heading_cn(doc, "P18  实际对照、消融与方案优势", 1)
    add_table(
        doc,
        ["实际比较对象", "保持相同的数据需求", "采用的评价方法", "本作品结果", "对照结果", "结论与代价"],
        [
            [
                "仅规划 vs 真实检索",
                "同一 HER2+/PIK3CA/治疗响应问题",
                "质量门、行数、是否访问外部库",
                "真实检索：848×46，来源 PASS，结局 REVIEW",
                "仅规划：0 行，不访问外部库",
                "检索能拿到真实宽表，但不能单凭行数宣称可分析",
            ],
            [
                "有质量门 vs 「有表即发布」",
                "同上",
                "分析数据集是否放行",
                "质量门拦截，Analysis Dataset=0",
                "若无门控，848 行可能被误当作响应队列",
                "门控增加「看起来不完整」的观感，避免错误科研使用",
            ],
            [
                "确定性规划 vs 千问 Function Calling",
                "同上",
                "used_qwen、工具选择、摘要是否虚构",
                "代表案例为确定性规划",
                "待实测（需百炼会话）",
                "未完成前不写优劣分数",
            ],
            [
                "单源 METABRIC vs 多源强行合并响应",
                "同上",
                "实体对齐、response_domain",
                "不跨库并患者",
                "手工把 GEO 响应贴到 METABRIC 患者",
                "后者不可追溯，本系统拒绝",
            ],
        ],
    )
    add_para(doc, "科研数据适用性方面：把适用性写成 Gate 4 与漏斗第三段，避免「库找到了 = 能做该分析」。数据处理方法方面：PICO → 受控工具 → 宽表 → 身份命名空间 → 四层门 → 缺口回查；清洗与原始值并存。结果质量和复用方面：可导出 Metadata/质量报告/CSV/Parquet/Excel；字段字典中文；来源可点回官方 URL。")
    add_para(doc, "没有改善的部分及增加的成本：冻结 SDTI 与外部 Benchmark（BEIR/Valentine 等）尚未同口径实测，不能填写分数；多源并不能自动补全「分子在 A 库、响应在 B 库且患者不可对齐」的研究；真实检索有网络与时延成本；质量门导致部分任务没有「漂亮的最终表」。")

    add_heading_cn(doc, "交给成员 A 的 P19 / P20 素材", 1)
    add_para(doc, "总体表现可引用：任务范围是乳腺癌科研问题 → 公开库查找解析整合。代表案例处理 848 条样本级记录、46 个字段、44 个已登记来源。主要质量结果：来源可追溯 PASS，实体 MATCH，科研适用性 REVIEW。能找到原始出处（官方 URL + accession + 本地缓存），但当前输出不适合直接做治疗响应建模。")
    add_para(doc, "必须人工审核：IHC 2+、跨库是否同一患者、结局域是否匹配、正式评测成绩。尚不能支持：无 Gold Set 的官方 SDTI 宣称；GTEx/OncoKB 等未接入库的假数据；把细胞系药敏当成患者疗效。")
    add_para(doc, "复现入口：仓库前端 http://127.0.0.1:8002/ ；POST /api/research/task ；GET /api/task/status/{id} ；导出 /export/{format} ；千问 POST /api/agent/qwen-sessions。源代码与依赖见 README_START_HERE.md。")

    add_heading_cn(doc, "附录  医学安全边界（本作品硬约束）", 1)
    for item in [
        "HER2 IHC 2+ 不得直接自动判为 HER2 Positive。",
        "ERBB2 CNA amplification 不等同 HER2 IHC positive。",
        "低置信度患者/样本关联进入 unresolved/review。",
        "高权威来源不可解释冲突不得自动选边。",
        "细胞系 AUC/IC50 与患者 pCR/response 必须用 response_domain 区分。",
        "不得通过硬编码 benchmark 答案通过测试，不得生成虚假系统成绩。",
    ]:
        add_para(doc, "• " + item, first_line=False, space_after=4)

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.paragraph_format.space_before = Pt(18)
    r = footer.add_run("— 完 —")
    set_run_font(r, size=11, color=MUTED)

    out1 = ROOT / "徐士诚_方向1A_P5-P18_报告.docx"
    out2 = DESKTOP / "徐士诚_方向1A_P5-P18_报告.docx"
    doc.save(out1)
    doc.save(out2)
    print(out1)
    print(out2)


if __name__ == "__main__":
    build()
